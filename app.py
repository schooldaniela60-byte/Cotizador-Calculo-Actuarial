"""
app.py — Interfaz Streamlit del PIA Calculo Actuarial
Requiere: TABLALX.txt y main.py en la misma carpeta
Correr con: python -m streamlit run app.py
"""

import streamlit as st
import pandas as pd
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io
from docx.shared import Inches
# Inicializar lista de cotizaciones en memoria
if "cotizaciones" not in st.session_state:
    st.session_state.cotizaciones = []

if "num_cotizaciones" not in st.session_state:
    st.session_state.num_cotizaciones = 0

# CONFIGURACIÓN DE PÁGINA

st.set_page_config(page_title="PIA Calculo Actuarial", layout="wide")
st.markdown("""
<style>

/* ── FONDO GENERAL ─────── */
.stApp {
    background-color: #FFFFFF !important;
    background-image: none !important;
}

/* ── ENCABEZADO ────────── */
h1 {
    color: #0A2F6E !important;
    font-weight: 800 !important;
    font-size: 2.2rem !important;
}

h2, h3 {
    color: #1565C0 !important;
    font-weight: 700 !important;
}

/* ── CAPTION  */
div[data-testid="stCaptionContainer"] p {
    color: #5C7A9F !important;
}

/* ── LÍNEAS DIVISORAS ──── */
hr {
    border: none !important;
    border-top: 1.5px solid #BBDEFB !important;
    margin: 1.5rem 0 !important;
}

/* ── LABELS ────────────── */
label[data-testid="stWidgetLabel"] p,
div[data-testid="stWidgetLabel"] p {
    color: #0A2F6E !important;
    font-weight: 600 !important;
    font-size: 13px !important;
}

/* ── INPUTS ────────────── */
input, textarea {
    color: #0A2F6E !important;
    font-weight: 600 !important;
    caret-color: #0A2F6E !important;
    border: 1.5px solid #90CAF9 !important;
    border-radius: 8px !important;
}

input:focus, textarea:focus {
    border-color: #1565C0 !important;
    box-shadow: 0 0 0 2px rgba(21,101,192,0.15) !important;
}

/* ── RADIO BUTTONS ─────── */
input[type="radio"] {
    accent-color: #1565C0 !important;
}

/* ── CHECKBOX ──────────── */
input[type="checkbox"] {
    accent-color: #1565C0 !important;
}

/* ── SELECTBOX ─────────── */
div[data-baseweb="select"] > div {
    border: 1.5px solid #90CAF9 !important;
    border-radius: 8px !important;
    background-color: #FFFFFF !important;
    cursor: default !important;
}

/* APUNTADO ULTRA-ESTRICTO PARA LOGRAR LA NEGRITA EN EL VALOR FIJADO */
div[data-baseweb="select"] [data-testid="stSelectboxSelectedValue"],
div[data-baseweb="select"] [data-testid="stSelectboxSelectedValue"] *,
div[data-baseweb="select"] [role="button"] p,
div[data-baseweb="select"] [role="button"] span {
    color: #0A2F6E !important;
    font-weight: 800 !important; /* Subido a 800 (Extra Bold) para asegurar el cambio visual */
    font-family: sans-serif !important; 
}
/* Texto seleccionado en selectbox en negritas */
div[data-baseweb="select"] [data-testid="stMarkdownContainer"] p,
div[data-baseweb="select"] span[class*="placeholder"],
div[data-baseweb="select"] span,
div[data-baseweb="single-select"] span,
div[data-baseweb="select"] [role="combobox"] span {
    font-weight: 700 !important;
    color: #0A2F6E !important;
}
/* DESINTEGRACIÓN TOTAL DEL INPUT INTERNO (EVITA CURSORES Y BORDES) */
div[data-baseweb="select"] input {
    position: absolute !important;
    width: 0 !important;
    height: 0 !important;
    padding: 0 !important;
    margin: 0 !important;
    border: none !important;
    outline: none !important;
    background: transparent !important;
    caret-color: transparent !important;
    color: transparent !important;
    opacity: 0 !important;
    pointer-events: none !important;
    user-select: none !important;
}

/* Quitar cualquier borde de enfoque residual */
div[data-baseweb="select"] > div:focus-within,
div[data-baseweb="select"] div:focus,
div[data-baseweb="select"] [role="combobox"]:focus {
    border-color: #1565C0 !important;
    box-shadow: 0 0 0 2px rgba(21,101,192,0.15) !important;
    outline: none !important;
}

/* ── MENÚ DESPLEGABLE ──── */
ul[data-baseweb="menu"],
div[data-baseweb="popover"] {
    background-color: #FFFFFF !important;
    border: 1px solid #90CAF9 !important;
    border-radius: 10px !important;
}

li[role="option"] {
    background-color: #FFFFFF !important;
    color: #0A2F6E !important;
    cursor: pointer !important;
    font-weight: normal !important; /* <--- Opciones del menú con grosor normal */
}

li[role="option"]:hover {
    background-color: #E3F2FD !important;
    color: #0A2F6E !important;
}

li[role="option"][aria-selected="true"] {
    background-color: #BBDEFB !important;
    color: #0A2F6E !important;
    font-weight: normal !important; /* Mantiene normal la opción activa en la lista */
}
/* ── PESTAÑAS ──────────── */
div[data-baseweb="tab-list"] {
    background-color: #E3F2FD !important;
    border-radius: 10px !important;
    padding: 4px !important;
    gap: 4px !important;
}

button[data-baseweb="tab"] {
    color: #1565C0 !important;
    font-weight: 600 !important;
    border-radius: 8px !important;
    background-color: transparent !important;
    padding: 8px 20px !important;
}

button[data-baseweb="tab"][aria-selected="true"] {
    color: #FFFFFF !important;
    font-weight: 700 !important;
    background: linear-gradient(90deg, #1565C0, #0A2F6E) !important;
    border-radius: 8px !important;
    border-bottom: none !important;
}

button[data-baseweb="tab"]:hover {
    background-color: rgba(21,101,192,0.12) !important;
    color: #0A2F6E !important;
}

/* ── BOTÓN CALCULAR ────── */
div[data-testid="stButton"] > button {
    background: linear-gradient(90deg, #1976D2, #0A2F6E) !important;
    color: #FFFFFF !important;
    font-weight: 700 !important;
    font-size: 15px !important;
    border: none !important;
    border-radius: 10px !important;
    padding: 12px 24px !important;
    transition: all 0.2s ease !important;
}

div[data-testid="stButton"] > button:hover {
    background: linear-gradient(90deg, #1E88E5, #1565C0) !important;
    transform: translateY(-1px) !important;
    box-shadow: 0 4px 16px rgba(21,101,192,0.35) !important;
}

/* ── RESULTADOS ────────── */
[data-testid="stMetric"],
[data-testid="metric-container"] {
    background: linear-gradient(135deg, #E3F2FD, #FFFFFF) !important;
    border: 1.5px solid #90CAF9 !important;
    border-radius: 14px !important;
    padding: 20px !important;
    box-shadow: 0 2px 8px rgba(21,101,192,0.1) !important;
}

[data-testid="stMetricLabel"] p,
[data-testid="metric-container"] label {
    color: #1565C0 !important;
    font-weight: 600 !important;
    font-size: 12px !important;
    text-transform: uppercase !important;
    letter-spacing: 0.05em !important;
}

[data-testid="stMetricValue"] {
    color: #0A2F6E !important;
    font-weight: 800 !important;
    font-size: 26px !important;
}

/* ── SIDEBAR  */
section[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #E3F2FD, #FFFFFF) !important;
    border-right: 1.5px solid #90CAF9 !important;
}

/* ── NÚMERO INPUT ──────── */
div[data-testid="stNumberInput"] input {
    color: #0A2F6E !important;
    font-weight: 600 !important;
}
/* Negrita selectbox — fuerza con atributo de estilo */
[data-baseweb="select"] [aria-selected="true"],
[data-baseweb="select"] [aria-selected],
[data-baseweb="select"] > div > div > div {
    font-weight: 560 !important;
    color: #0A2F6E !important;
}
</style>
""", unsafe_allow_html=True)



# CARGAR TABLA ACTUARIAL


@st.cache_data
def cargar_tabla():
    ruta = os.path.join(os.path.dirname(os.path.abspath(__file__)), "TABLALX.txt")
    COLS = ["Edad", "qx", "px", "lx", "dx", "Cx", "Dx", "Mx", "Nx"]
    tabla = pd.read_csv(ruta, sep=r"\s*\|\s*", skiprows=2, names=COLS, engine="python")
    tabla = tabla.apply(lambda x: x.astype(str).str.strip())
    for col in COLS:
        tabla[col] = pd.to_numeric(tabla[col], errors="coerce")
    tabla = tabla.dropna()
    return {
        "edades": tabla["Edad"].astype(int).tolist(),
        "qx": tabla["qx"].tolist(),
        "px": tabla["px"].tolist(),
        "lx": tabla["lx"].tolist(),
        "dx": tabla["dx"].tolist(),
        "Cx": tabla["Cx"].tolist(),
        "Dx": tabla["Dx"].tolist(),
        "Mx": tabla["Mx"].tolist(),
        "Nx": tabla["Nx"].tolist(),
    }

T = cargar_tabla()

def idx(edad):
    return T["edades"].index(int(edad))


# ENCABEZADO

def formatear_monto(key):
    val = st.session_state[key].replace(",", "").replace("$", "").strip()
    try:
        numero = float(val)
        st.session_state[key] = f"{numero:,.2f}"
    except:
        pass

st.title("Cotizador Calculo Actuarial")
st.caption("CNSF 2000I · i = 5%")
st.session_state.num_cotizaciones = st.number_input(
    "¿Cuántas cotizaciones deseas realizar?",
    min_value=1, max_value=20, value=1, step=1
)

completadas = len(st.session_state.cotizaciones)
total = int(st.session_state.num_cotizaciones)
progreso = min(completadas / total, 1.0) if total > 0 else 0.0
st.progress(
    progreso,
    text=f"Cotización {completadas} de {total} completadas"
)
# ── Datos del asegurado (FUERA de cualquier tab) ──
st.subheader("Datos del asegurado")
col1, col2, col3, col4 = st.columns(4)

with col1:
    nombre = st.text_input("Nombre completo")
with col2:
    edad_real = st.number_input("Edad", min_value=18, max_value=70, value=35)
with col3:
    sexo = st.selectbox("Sexo", ["Hombre", "Mujer"])
with col4:
    fumador = st.radio("Fumador/a", ["Sí", "No"])

descuento = 0
if sexo == "Mujer":
    descuento += 2
if fumador == "No":
    descuento += 2
edad_act = max(int(edad_real) - descuento, 18)

st.divider()

# ── Pestañas 
tab1, tab2, tab3 = st.tabs(["Seguros de vida", "Anualidades", "Tabla actuarial"])



#  — SEGUROS DE VIDA


with tab1:

    st.subheader("Seguro de vida")

    col1, col2 = st.columns(2)

    with col1:
        tipo_seguro = st.selectbox("Tipo de seguro", [
            "Temporal",
            "Diferido temporal",
            "Diferido vitalicio",
            "Dotal puro",
            "Dotal mixto",
            "Vitalicio"
        ])


        SA = st.number_input(
            "Suma asegurada ($)",
            min_value=0.0,
            value=1000000.0,
            placeholder="Ej: 1000000",
            step=1000000.0,
            format="%.2f"
        )
        st.caption(f"${SA:,.2f}")


    with col2:

        n, m, t = None, None, None

        if tipo_seguro in ["Temporal", "Diferido temporal"]:
            n = st.number_input("Temporalidad (n)", min_value=1, value=20)

        if tipo_seguro in ["Dotal puro", "Dotal mixto"]:
            n = st.number_input("Temporalidad (n) — máx. hasta edad 65", min_value=1, value=20)

        if tipo_seguro in ["Diferido temporal", "Diferido vitalicio", "Dotal puro", "Dotal mixto"]:
            m = st.number_input("Diferimiento (m)", min_value=0, value=0)

        if tipo_seguro in ["Temporal", "Diferido temporal", "Diferido vitalicio",
                           "Dotal puro", "Dotal mixto", "Vitalicio"]:
            pagos_limitados = st.checkbox("¿Pagos limitados?", key="chk_pl")
            if pagos_limitados:
                t = st.number_input("Años de pago (t)", min_value=1, value=10, key="t_input")

    st.divider()

    if st.button("Calcular", type="primary", use_container_width=True, key="btn_seguros"):
        try:

            x = edad_act
            n = int(n) if n is not None else None
            m = int(m) if m is not None else 0
            t = int(t) if t is not None else None

            # ── Temporal ─
            if tipo_seguro == "Temporal":
                t_calc = t if t is not None else n
                if x + n > 100:
                    st.error("La edad final excede el límite de la tabla (100).")
                else:
                    i  = idx(x);  j  = idx(x + n);  it = idx(x + t_calc)
                    A   = (T["Mx"][i] - T["Mx"][j]) / T["Dx"][i]
                    PNU = SA * A
                    ann  = (T["Nx"][i] - T["Nx"][j])  / T["Dx"][i]
                    PNN  = PNU / ann
                    annL = (T["Nx"][i] - T["Nx"][it]) / T["Dx"][i]
                    PNNL = PNU / annL

                    st.subheader(f"Cotización — {nombre if nombre else 'Asegurado'}")
                    st.caption(f"Seguro Temporal · Edad actuarial: {x} años · Temporalidad: {n} años")
                    c1, c2, c3, c4 = st.columns(4)
                    c1.metric("Valores Conmutados",        f"{A:.6f}")
                    c2.metric("PNU",             f"${PNU:,.2f}")
                    c3.metric("PNN ",        f"${PNN:,.2f}")
                    c4.metric("PNN pag. limit.",  f"${PNNL:,.2f}")

            # ── Diferido temporal 
            elif tipo_seguro == "Diferido temporal":
                if x + m + n > 100:
                    st.error("La edad final excede el límite de la tabla (100).")
                else:
                    i   = idx(x);  im  = idx(x + m);  imn = idx(x + m + n)
                    t_calc = t if t is not None else n
                    it  = idx(x + t_calc)
                    A   = (T["Mx"][im] - T["Mx"][imn]) / T["Dx"][i]
                    PNU = SA * A
                    ann  = (T["Nx"][i] - T["Nx"][idx(x + n)]) / T["Dx"][i]
                    PNN  = PNU / ann
                    annL = (T["Nx"][i] - T["Nx"][it]) / T["Dx"][i]
                    PNNL = PNU / annL

                    st.subheader(f"Cotización — {nombre if nombre else 'Asegurado'}")
                    st.caption(f"Seguro Diferido Temporal · Edad actuarial: {x} · Diferimiento: {m} · Temporalidad: {n}")
                    c1, c2, c3, c4 = st.columns(4)
                    c1.metric("Valores Conmutados", f"{A:.6f}")
                    c2.metric("PNU",                f"${PNU:,.2f}")
                    c3.metric("PNN",           f"${PNN:,.2f}")
                    c4.metric("PNN pag. limit.",     f"${PNNL:,.2f}")

            # ── Diferido vitalicio 
            elif tipo_seguro == "Diferido vitalicio":
                if x + m > 100:
                    st.error("La edad final excede el límite de la tabla (100).")
                else:
                    i   = idx(x);  im  = idx(x + m)
                    t_calc = t if t is not None else None
                    A   = T["Mx"][im] / T["Dx"][i]
                    PNU = SA * A
                    ann = T["Nx"][i] / T["Dx"][i]
                    PNN = PNU / ann

                    st.subheader(f"Cotización — {nombre if nombre else 'Asegurado'}")
                    st.caption(f"Seguro Diferido Vitalicio · Edad actuarial: {x} · Diferimiento: {m}")

                    if t_calc is not None:
                        it   = idx(x + t_calc)
                        annL = (T["Nx"][i] - T["Nx"][it]) / T["Dx"][i]
                        PNNL = PNU / annL
                        c1, c2, c3, c4 = st.columns(4)
                        c1.metric("Valores Conmutados", f"{A:.6f}")
                        c2.metric("PNU",                f"${PNU:,.2f}")
                        c3.metric("PNN ",      f"${PNN:,.2f}")
                        c4.metric("PNN pag. limit.",    f"${PNNL:,.2f}")
                    else:
                        c1, c2, c3 = st.columns(3)
                        c1.metric("Valores Conmutados", f"{A:.6f}")
                        c2.metric("PNU",                f"${PNU:,.2f}")
                        c3.metric("PNN",      f"${PNN:,.2f}")

            # ── Dotal puro ─
            elif tipo_seguro == "Dotal puro":
                if x + n > 65:
                    st.error(f"La edad final no puede exceder 65 años. Edad actuarial {x} + temporalidad {n} = {x+n}.")
                elif x + m + n > 100:
                    st.error("La edad final excede el límite de la tabla (100).")
                else:
                    i   = idx(x); imn = idx(x + m + n)
                    A   = T["Dx"][imn] / T["Dx"][i]
                    PNU = SA * A
                    ann = (T["Nx"][i] - T["Nx"][idx(x + n)]) / T["Dx"][i]
                    PNN = PNU / ann

                    st.subheader(f"Cotización — {nombre if nombre else 'Asegurado'}")

                    if t is not None:
                        it   = idx(x + t)
                        annL = (T["Nx"][i] - T["Nx"][it]) / T["Dx"][i]
                        PNNL = PNU / annL
                        st.caption(f"Dotal Puro · Edad actuarial: {x} · Temporalidad: {n} · Pagos limitados: {t} años")
                        c1, c2, c3, c4 = st.columns(4)
                        c1.metric("Factor A",       f"{A:.6f}")
                        c2.metric("PNU",            f"${PNU:,.2f}")
                        c3.metric("PNN",      f"${PNN:,.2f}")
                        c4.metric("PNN pag. limit.",f"${PNNL:,.2f}")
                    else:
                        st.caption(f"Dotal Puro · Edad actuarial: {x} · Temporalidad: {n}")
                        c1, c2, c3 = st.columns(3)
                        c1.metric("Factor A",  f"{A:.6f}")
                        c2.metric("PNU",       f"${PNU:,.2f}")
                        c3.metric("PNN ", f"${PNN:,.2f}")

            # ── Dotal mixto 

            elif tipo_seguro == "Dotal mixto":
                if x + n > 65:
                    st.error(f"La edad final no puede exceder 65 años. Edad actuarial {x} + temporalidad {n} = {x+n}.")
                elif x + m + n > 100:
                    st.error("La edad final excede el límite de la tabla (100).")
                else:
                    i   = idx(x); im = idx(x + m); imn = idx(x + m + n)
                    A   = (T["Mx"][im] - T["Mx"][imn] + T["Dx"][imn]) / T["Dx"][i]
                    PNU = SA * A
                    ann = (T["Nx"][i] - T["Nx"][idx(x + n)]) / T["Dx"][i]
                    PNN = PNU / ann

                    st.subheader(f"Cotización — {nombre if nombre else 'Asegurado'}")

                    if t is not None:
                        it   = idx(x + t)
                        annL = (T["Nx"][i] - T["Nx"][it]) / T["Dx"][i]
                        PNNL = PNU / annL
                        st.caption(f"Dotal Mixto · Edad actuarial: {x} · Temporalidad: {n} · Pagos limitados: {t} años")
                        c1, c2, c3, c4 = st.columns(4)
                        c1.metric("Factor A",       f"{A:.6f}")
                        c2.metric("PNU",            f"${PNU:,.2f}")
                        c3.metric("PNN ",      f"${PNN:,.2f}")
                        c4.metric("PNN pag. limit.",f"${PNNL:,.2f}")
                    else:
                        st.caption(f"Dotal Mixto · Edad actuarial: {x} · Temporalidad: {n}")
                        c1, c2, c3 = st.columns(3)
                        c1.metric("Factor A",  f"{A:.6f}")
                        c2.metric("PNU",       f"${PNU:,.2f}")
                        c3.metric("PNN ", f"${PNN:,.2f}")

            # ── Vitalicio 
            elif tipo_seguro == "Vitalicio":
                i   = idx(x)
                A   = T["Mx"][i] / T["Dx"][i]
                PNU = SA * A
                ann = T["Nx"][i] / T["Dx"][i]
                PNN = PNU / ann

                st.subheader(f"Cotización — {nombre if nombre else 'Asegurado'}")

                if t is not None:
                    it   = idx(x + t)
                    annL = (T["Nx"][i] - T["Nx"][it]) / T["Dx"][i]
                    PNNL = PNU / annL
                    st.caption(f"Vitalicio · Edad actuarial: {x} · Pagos limitados: {t} años")
                    c1, c2, c3, c4 = st.columns(4)
                    c1.metric("Factor A",       f"{A:.6f}")
                    c2.metric("PNU",            f"${PNU:,.2f}")
                    c3.metric("PNN ",  f"${PNN:,.2f}")
                    c4.metric("PNN pag. limit.",f"${PNNL:,.2f}")
                else:
                    st.caption(f"Vitalicio · Edad actuarial: {x}")
                    c1, c2, c3 = st.columns(3)
                    c1.metric("Factor A",     f"{A:.6f}")
                    c2.metric("PNU",          f"${PNU:,.2f}")
                    c3.metric("PNN ",f"${PNN:,.2f}")

            st.session_state.cotizaciones.append({
                "fecha":        datetime.now().strftime("%d/%m/%Y %H:%M"),
                "nombre":       nombre if nombre else "Asegurado",
                "edad_real":    edad_real,
                "edad_act":     edad_act,
                "sexo":         sexo,
                "fumador":      fumador,
                "tipo":         tipo_seguro,        # o tipo_ann para anualidades
                "parametros": {
                    "Temporalidad": n if n else "—",
                    "Diferimiento": m if m else "—",
                    "Núm. pagos":   t if t else "—",
                },
                "resultados": {
                    "Valores Conmutados": f"{A:.6f}",
                    "PNU":      f"${PNU:,.2f}",
                    "PNN":      f"${PNN:,.2f}",
                }
            })
            st.success(f"Cotización {len(st.session_state.cotizaciones)} guardada.")
        except ValueError as e:
            st.error(f"Error en los datos: {e}")
        except Exception as e:
            st.error(f"Error inesperado: {e}")
        # Guardar en memoria


#  — ANUALIDADES
# =====================================================================
# REEMPLAZA TU BLOQUE DE "WITH TAB2:" POR ESTE COMPLETAMENTE CORREGIDO
# =====================================================================

with tab2:
    st.subheader("Anualidades")

    col1, col2 = st.columns(2)

    with col1:
        tipo_ann = st.selectbox("Tipo de anualidad", [
            "Temporal", "Vitalicia", "Diferida temporal", "Diferida vitalicia"
        ], key="sb_tipo_ann_unique")
        
        renta = st.number_input(
            "Renta periódica ($)",
            min_value=0.0,
            value=10000.0,
            placeholder="Ej: 30000",
            step=10000.0,
            format="%.2f",
            key="num_renta_unique"
        )
        st.caption(f"${renta:,.2f}")
        modalidad = st.radio("Modalidad", ["Anticipada", "Vencida"], key="rd_modalidad_unique")

    with col2:
        n_ann, m_ann = None, None

        if tipo_ann in ["Temporal", "Diferida temporal"]:
            n_ann = st.number_input("Temporalidad (m)", min_value=1, value=15, key="num_n_ann_unique")

        if tipo_ann in ["Diferida temporal", "Diferida vitalicia"]:
            m_ann = st.number_input("Diferimiento (n)", min_value=0, value=5, key="num_m_ann_unique")

    st.divider()

    # 1. EJECUCIÓN DEL CÁLCULO AL PRESIONAR EL BOTÓN
    if st.button("Calcular Anualidad", type="primary", use_container_width=True, key="btn_anualidades_unique"):
        try:
            x = edad_act
            anticipada = (modalidad == "Anticipada")
            n_val = int(n_ann) if n_ann is not None else None
            m_val = int(m_ann) if m_ann is not None else 0

            i_x = idx(x)
            Dx = T["Dx"][i_x]
            factor = 0.0
            tipo_str = ""

            # ── Temporal ──────────────────────────────────────────────────
            if tipo_ann == "Temporal":
                if anticipada:
                    i_xm = idx(x + n_val)
                    factor = (T["Nx"][i_x] - T["Nx"][i_xm]) / Dx
                    tipo_str = "Temporal Anticipada"
                else:
                    i_x1 = idx(x + 1)
                    i_xm1 = idx(x + n_val + 1)
                    factor = (T["Nx"][i_x1] - T["Nx"][i_xm1]) / Dx
                    tipo_str = "Temporal Vencida"

            # ── Vitalicia ─────────────────────────────────────────────────
            elif tipo_ann == "Vitalicia":
                if anticipada:
                    factor = (T["Nx"][i_x]) / Dx
                    tipo_str = "Vitalicia Anticipada"
                else:
                    i_x1 = idx(x + 1)
                    factor = T["Nx"][i_x1] / Dx
                    tipo_str = "Vitalicia Vencida"

            # ── Diferida temporal ─────────────────────────────────────────
            elif tipo_ann == "Diferida temporal":
                if anticipada:
                    i_n = idx(x + m_val)
                    i_nm = idx(x + m_val + n_val)
                    factor = (T["Nx"][i_n] - T["Nx"][i_nm]) / Dx
                    tipo_str = "Diferida Temporal Anticipada"
                else:
                    i_n1 = idx(x + m_val + 1)
                    i_nm1 = idx(x + m_val + n_val + 1)
                    factor = (T["Nx"][i_n1] - T["Nx"][i_nm1]) / Dx
                    tipo_str = "Diferida Temporal Vencida"

            # ── Diferida vitalicia ────────────────────────────────────────
            elif tipo_ann == "Diferida vitalicia":
                if anticipada:
                    i_m = idx(x + m_val)
                    factor = (T["Nx"][i_m]) / Dx
                    tipo_str = "Diferida Vitalicia Anticipada"
                else:
                    i_m1 = idx(x + m_val + 1)
                    factor = T["Nx"][i_m1] / Dx
                    tipo_str = "Diferida Vitalicia Vencida"

            prima = renta * factor

            # PERSISTENCIA: Guardamos el cálculo actual en una variable específica para mostrar en pantalla
            st.session_state["ultimo_calculo_ann"] = {
                "tipo_str": tipo_str,
                "nombre": nombre if nombre else "Asegurado",
                "edad_act": x,
                "factor": factor,
                "renta": renta,
                "prima": prima
            }

            # Guardamos en el historial global (para el Excel/Word)
            st.session_state.cotizaciones.append({
                "fecha": datetime.now().strftime("%d/%m/%Y %H:%M"),
                "nombre": nombre if nombre else "Asegurado",
                "edad_real": edad_real,
                "edad_act": edad_act,
                "sexo": sexo,
                "fumador": fumador,
                "tipo": tipo_ann,
                "parametros": {
                    "Renta periódica": f"${renta:,.2f}",
                    "Modalidad": modalidad,
                    "Temporalidad": f"{n_val} años" if n_val else "—",
                    "Diferimiento": f"{m_val} años" if m_val else "—",
                },
                "resultados": {
                    "VALORES CONMUTADOS": f"{factor:.4f}",
                    "Prima única": f"${prima:,.2f}",
                }
            })
            st.success(f"Cotización {len(st.session_state.cotizaciones)} guardada.")
            
        except ValueError as e:
            st.error(f"Error en los datos: {e}")
        except Exception as e:
            st.error(f"Error inesperado: {e}")

    # 2. BLOQUE DE IMPRESIÓN EN PANTALLA (Fuera del botón para que no desaparezca)
    if "ultimo_calculo_ann" in st.session_state and st.session_state["ultimo_calculo_ann"] is not None:
        res = st.session_state["ultimo_calculo_ann"]
        
        st.markdown("---")
        st.subheader(f"Cotización — {res['nombre']}")
        st.caption(f"{res['tipo_str']} · Edad actuarial: {res['edad_act']} años")
        
        # Crear las tres tarjetas de resultados en columnas
        c1, c2, c3 = st.columns(3)
        c1.metric(label="VALORES CONMUTADOS", value=f"{res['factor']:.5f}")
        c2.metric(label="RENTA PERIÓDICA", value=f"${res['renta']:,.2f}")
        c3.metric(label="VALOR PRESENTE ANUALIDAD", value=f"${res['prima']:,.2f}")

with tab3:

    st.subheader("Tabla de Mortalidad")
    st.caption("CNSF 2000I · i = 5% · Radix = 1,000,000")

    df = pd.DataFrame({
        "Edad": T["edades"],
        "qx":   [f"{v:.6f}" for v in T["qx"]],
        "px":   [f"{v:.6f}" for v in T["px"]],
        "lx":   [f"{v:,.0f}" for v in T["lx"]],
        "dx":   [f"{v:,.0f}" for v in T["dx"]],
        "Cx":   [f"{v:,.3f}" for v in T["Cx"]],
        "Dx":   [f"{v:,.3f}" for v in T["Dx"]],
        "Mx":   [f"{v:,.3f}" for v in T["Mx"]],
        "Nx":   [f"{v:,.3f}" for v in T["Nx"]],
    })

    st.dataframe(df, use_container_width=True, hide_index=True, height=600)
# ── EXPORTAR WORD 
if len(st.session_state.cotizaciones) > 0:
    st.divider()

    col_info, col_limpiar = st.columns([3, 1])
    with col_info:
        st.markdown(f"**{len(st.session_state.cotizaciones)}** cotización(es) guardada(s)")
    with col_limpiar:
        if st.button("Limpiar memoria", key="limpiar"):
            st.session_state.cotizaciones = []
            st.rerun()

    def generar_word(cotizaciones):
        doc = Document()

        AZUL_OSCURO = RGBColor(0x0A, 0x2F, 0x6E)
        AZUL_MEDIO  = RGBColor(0x15, 0x65, 0xC0)

        # ── Portada 
        doc.add_paragraph("")
        doc.add_paragraph("")

        # Título principal
        titulo = doc.add_heading("Cotizador Calculo Actuarial", 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        titulo.runs[0].font.color.rgb = AZUL_OSCURO
        titulo.runs[0].font.size = Pt(28)

        # Línea decorativa
        linea = doc.add_paragraph("─" * 50)
        linea.alignment = WD_ALIGN_PARAGRAPH.CENTER
        linea.runs[0].font.color.rgb = AZUL_MEDIO

        # Subtítulo
        sub = doc.add_paragraph("CNSF 2000I · Tasa de interés: 5%")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.runs[0].font.color.rgb = AZUL_MEDIO
        sub.runs[0].font.size = Pt(13)

        doc.add_paragraph("")

        # Fecha
        fecha_doc = doc.add_paragraph(
            f"Generado el: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
        )
        fecha_doc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fecha_doc.runs[0].font.color.rgb = AZUL_OSCURO
        fecha_doc.runs[0].font.size = Pt(11)

        doc.add_paragraph("")
        doc.add_paragraph("")

        # ── Tabla de contenido ────
        h_toc = doc.add_heading("Tabla de contenido", level=1)
        h_toc.runs[0].font.color.rgb = AZUL_OSCURO

        toc = doc.add_table(rows=1, cols=4)
        toc.style = "Table Grid"

        # Encabezados
        headers = ["#", "Nombre", "Tipo de producto", "Fecha"]
        for j, header in enumerate(headers):
            cell = toc.rows[0].cells[j]
            cell.text = header
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run.font.color.rgb = AZUL_OSCURO

        # Filas por cotización
        for i, c in enumerate(cotizaciones):
            row = toc.add_row().cells
            row[0].text = str(i + 1)
            row[1].text = c["nombre"]
            row[2].text = c["tipo"]
            row[3].text = c["fecha"]

        # Ancho de columnas

        toc.columns[0].width = Inches(0.4)
        toc.columns[1].width = Inches(1.8)
        toc.columns[2].width = Inches(2.5)
        toc.columns[3].width = Inches(1.5)

        doc.add_paragraph("")

        # Total
        total_p = doc.add_paragraph(
            f"Total de cotizaciones: {len(cotizaciones)}"
        )
        total_p.runs[0].font.bold = True
        total_p.runs[0].font.color.rgb = AZUL_OSCURO

        doc.add_page_break()

        # ── Una sección por cotización ────────────────
        for i, c in enumerate(cotizaciones):

            # Título cotización
            h1 = doc.add_heading(f"Cotización #{i+1}", level=1)
            h1.runs[0].font.color.rgb = AZUL_OSCURO

            doc.add_paragraph(f"Fecha: {c['fecha']}")
            doc.add_paragraph("")

            # Datos del asegurado
            h2 = doc.add_heading("Datos del asegurado", level=2)
            h2.runs[0].font.color.rgb = AZUL_MEDIO

            t1 = doc.add_table(rows=1, cols=2)
            t1.style = "Table Grid"
            t1.rows[0].cells[0].text = "Campo"
            t1.rows[0].cells[1].text = "Valor"
            for cell in t1.rows[0].cells:
                run = cell.paragraphs[0].runs[0]
                run.font.bold = True
                run.font.color.rgb = AZUL_OSCURO

            for campo, valor in [
                ("Nombre",         c["nombre"]),
                ("Edad real",      f"{c['edad_real']} años"),
                ("Edad actuarial", f"{c['edad_act']} años"),
                ("Sexo",           c["sexo"]),
                ("Fumador/a",      c["fumador"]),
            ]:
                row = t1.add_row().cells
                row[0].text = campo
                row[1].text = str(valor)

            doc.add_paragraph("")

            # Tipo y parámetros
            h3 = doc.add_heading("Producto cotizado", level=2)
            h3.runs[0].font.color.rgb = AZUL_MEDIO

            doc.add_paragraph(f"Tipo: {c['tipo']}")

            t2 = doc.add_table(rows=1, cols=2)
            t2.style = "Table Grid"
            t2.rows[0].cells[0].text = "Parámetro"
            t2.rows[0].cells[1].text = "Valor"
            for cell in t2.rows[0].cells:
                run = cell.paragraphs[0].runs[0]
                run.font.bold = True
                run.font.color.rgb = AZUL_OSCURO

            for param, valor in c["parametros"].items():
                row = t2.add_row().cells
                row[0].text = param
                row[1].text = str(valor)

            doc.add_paragraph("")

            # Resultados
            h4 = doc.add_heading("Resultados", level=2)
            h4.runs[0].font.color.rgb = AZUL_MEDIO

            t3 = doc.add_table(rows=1, cols=2)
            t3.style = "Table Grid"
            t3.rows[0].cells[0].text = "Concepto"
            t3.rows[0].cells[1].text = "Valor"
            for cell in t3.rows[0].cells:
                run = cell.paragraphs[0].runs[0]
                run.font.bold = True
                run.font.color.rgb = AZUL_OSCURO

            for concepto, valor in c["resultados"].items():
                row = t3.add_row().cells
                row[0].text = concepto
                row[1].text = str(valor)

            if i < len(cotizaciones) - 1:
                doc.add_page_break()

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    buffer = generar_word(st.session_state.cotizaciones)

    st.download_button(
        label="Exportar todas las cotizaciones a Word",
        data=buffer,
        file_name=f"cotizaciones_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )